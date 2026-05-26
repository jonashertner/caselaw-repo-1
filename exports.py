"""
Practitioner workflow exports for OpenCaseLaw decisions.

Four formats are produced from the canonical decision row + reference
graph data already loaded by the MCP server:

  * `.docx`       — formatted decision (citation + Regeste + Sachverhalt
                    + Erwägungen + Dispositiv) ready to drop into a Word
                    file. Built with python-docx.
  * BibTeX        — `@misc{decision_id, title=..., url=..., year=...}`.
  * RIS           — TY-CASE record consumable by Zotero, EndNote, Mendeley.
  * Atom feed     — per-court "newest decisions" feed (top 50) for
                    practitioner subscriptions.

Each entry-point is a pure function that takes a decision dict (or a
court code) and returns `(body_bytes, media_type, suggested_filename)`
so the HTTP handlers in mcp_server.py just need to wrap the response.

The module deliberately has no external state — callers pass the
already-fetched row / row-list. Caching, DB connection, and rate-
limiting all stay in the calling layer.
"""
from __future__ import annotations

import io
import logging
import re
from datetime import datetime, timezone
from html import escape
from typing import Iterable

logger = logging.getLogger(__name__)

CANONICAL_BASE = "https://mcp.opencaselaw.ch"


# ── Common helpers ─────────────────────────────────────────────────

# XML 1.0 Char production:
#   #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD] | [#x10000-#x10FFFF]
# Anything outside that set makes python-docx (lxml) raise
#   ValueError: All strings must be XML compatible: Unicode or ASCII,
#   no NULL bytes or control characters
# Scraped Swiss court PDFs leak three classes of these:
#  (1) C0 controls — NULL/form-feed/etc. from PDF extraction
#  (2) lone surrogates (\ud800-\udfff) — broken UTF-16 reconstruction
#  (3) Unicode non-characters \ufffe / \uffff — observed in
#      e.g. vd_gerichte_FA13.055244 (4 \uffff in full_text) tripping
#      the L3 export-render QC gate on 2026-05-26.
# We strip all three at render-time so the export endpoint never 500s
# and the QC gate stays green on a clean corpus.
_XML_FORBIDDEN_CTRL = re.compile(
    "[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff\ufffe\uffff]"
)


def _xml_safe(s):
    """Return s with XML-incompatible characters removed (C0 controls,
    lone surrogates, and the \\ufffe/\\uffff non-characters)."""
    if not isinstance(s, str):
        return s
    return _XML_FORBIDDEN_CTRL.sub("", s)


def _xml_safe_decision(decision: dict) -> dict:
    """Return a shallow copy of `decision` with every string value
    sanitized of XML-forbidden control characters. Non-string values
    pass through unchanged."""
    return {k: _xml_safe(v) for k, v in decision.items()}


def _decision_url(decision_id: str) -> str:
    return f"{CANONICAL_BASE}/entscheid/{decision_id}"


def _safe_year(decision_date: str) -> str:
    if not decision_date:
        return ""
    m = re.match(r"(\d{4})", decision_date)
    return m.group(1) if m else ""


def _bib_key(decision_id: str) -> str:
    """Stable BibTeX cite-key from a decision_id (alnum only,
    no leading digit). Example bge_BGE_140_III_86 → bgeBGE140III86."""
    return re.sub(r"[^A-Za-z0-9]", "", decision_id) or "decision"


def _short_title(decision: dict, max_chars: int = 80) -> str:
    """Best-effort one-line title for a decision."""
    parts = []
    cs = decision.get("citation_string_de") or decision.get("citation_string")
    if cs:
        parts.append(cs)
    elif decision.get("docket_number"):
        court = decision.get("court_name") or decision.get("court") or ""
        parts.append(f"{court} {decision['docket_number']}".strip())
    else:
        parts.append(decision.get("decision_id", "decision"))
    if decision.get("title"):
        t = decision["title"].strip()
        if t and t.lower() not in (parts[0] or "").lower():
            parts.append(", " + t)
    out = " ".join(parts)
    if len(out) > max_chars:
        out = out[: max_chars - 1] + "..."
    return out


# ── BibTeX ─────────────────────────────────────────────────────────

def render_bibtex(decision: dict) -> tuple[bytes, str, str]:
    """Return one BibTeX `@misc` entry for a decision."""
    decision_id = decision.get("decision_id", "decision")
    key = _bib_key(decision_id)
    title = _short_title(decision)
    year = _safe_year(decision.get("decision_date") or "")
    court = decision.get("court_name") or decision.get("court") or ""
    url = _decision_url(decision_id)

    fields = [
        ("title",      title),
        ("author",     court),
        ("year",       year),
        ("howpublished", "OpenCaseLaw — opencaselaw.ch"),
        ("url",        url),
        ("note",       f"Docket: {decision.get('docket_number') or '—'}"),
    ]
    body_lines = [f"@misc{{{key},"]
    for k, v in fields:
        if not v:
            continue
        # BibTeX-escape: braces around the value protect spaces / case
        v_clean = str(v).replace("\\", " ").replace("{", "(").replace("}", ")")
        body_lines.append(f"  {k:<13} = {{{v_clean}}},")
    # Strip trailing comma on the last field
    body_lines[-1] = body_lines[-1].rstrip(",")
    body_lines.append("}")
    body = "\n".join(body_lines) + "\n"
    return body.encode("utf-8"), "application/x-bibtex; charset=utf-8", \
           f"{decision_id}.bib"


# ── RIS (Zotero / EndNote / Mendeley) ──────────────────────────────

def render_ris(decision: dict) -> tuple[bytes, str, str]:
    """RIS bibliographic record (type CASE for legal cases)."""
    decision_id = decision.get("decision_id", "decision")
    title = _short_title(decision, max_chars=200)
    year = _safe_year(decision.get("decision_date") or "")
    date_full = decision.get("decision_date") or ""
    court = decision.get("court_name") or decision.get("court") or ""
    docket = decision.get("docket_number") or ""
    url = _decision_url(decision_id)
    abstract = decision.get("regeste") or decision.get("abstract_de") \
               or decision.get("abstract_fr") or decision.get("abstract_it") \
               or ""

    lines = [
        "TY  - CASE",
        f"TI  - {title}",
        f"AU  - {court}",
    ]
    if year:
        lines.append(f"PY  - {year}")
    if date_full:
        # RIS DA format: YYYY/MM/DD/
        m = re.match(r"(\d{4})-(\d{2})-(\d{2})", date_full)
        if m:
            lines.append(f"DA  - {m.group(1)}/{m.group(2)}/{m.group(3)}/")
    if docket:
        lines.append(f"AN  - {docket}")           # accession number
    lang = (decision.get("language") or "").lower()
    if lang in ("de", "fr", "it"):
        lines.append(f"LA  - {lang}")
    if abstract:
        # RIS reads first 32k of AB; keep it tight
        lines.append(f"AB  - {abstract[:2000]}")
    lines.append(f"UR  - {url}")
    lines.append(f"ID  - {decision_id}")
    lines.append("ER  - ")
    body = "\r\n".join(lines) + "\r\n"
    return body.encode("utf-8"), "application/x-research-info-systems; charset=utf-8", \
           f"{decision_id}.ris"


# ── DOCX ───────────────────────────────────────────────────────────

# ── Text normalisation ────────────────────────────────────────────
#
# Critical for Word output: source text from the corpus is hard-wrapped
# (one source line per line, often 60–80 chars) because that's how
# courts publish decisions. python-docx's add_paragraph() converts every
# `\n` in the input into a `<w:br/>` hard line break — producing a
# Word document with 70+ broken lines per paragraph, where Word's own
# word-wrap should be doing the wrapping at the right margin.
#
# Fix: collapse single newlines into spaces (so the paragraph is one
# continuous string Word can re-flow). Caller splits on blank lines
# upstream when paragraph boundaries matter.

def _flow_paragraph(text: str) -> str:
    """Normalise a chunk of corpus text into a single re-flowable
    paragraph: collapse newlines and runs of whitespace into single
    spaces so Word's own word-wrap handles the line lengths."""
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\s*\n\s*", " ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _split_paragraphs(text: str) -> list[str]:
    """Split text on blank lines into logical paragraphs, then
    `_flow_paragraph` each one. Empty paragraphs are dropped."""
    if not text:
        return []
    parts = re.split(r"\n\s*\n+", text.replace("\r\n", "\n"))
    return [p for p in (_flow_paragraph(c) for c in parts) if p]


# Heuristic: regeste field for federal decisions concatenates DE / FR /
# IT versions back-to-back, each typically prefixed by "Regeste" /
# "Regeste" / "Regesto". Splitting at those boundaries gives one
# paragraph per language instead of a 600-char wall of mixed text.
_REGESTE_SPLIT_RE = re.compile(
    r"(?=\b(?:Regeste|Regesto)\b)",
    re.IGNORECASE,
)


def _split_regeste(regeste: str) -> list[str]:
    """Split a multilingual regeste into per-language paragraphs.
    Falls back to a single flowed paragraph when no language headers
    are detectable (most cantonal decisions)."""
    if not regeste or not regeste.strip():
        return []
    parts = [p for p in _REGESTE_SPLIT_RE.split(regeste) if p and p.strip()]
    if len(parts) <= 1:
        return [_flow_paragraph(regeste)]
    return [_flow_paragraph(p) for p in parts if _flow_paragraph(p)]


def render_docx(decision: dict, paragraphs: list[dict] | None = None) -> tuple[bytes, str, str]:
    """Build a Word document with the decision's structured content.

    `decision` is the row dict from `get_decision_by_id`. `paragraphs`
    is the structured-extraction sidecar (list of {e_number, text}) —
    optional; if absent we fall back to the raw full_text.

    Falls back to a plain-text `.txt` if python-docx is not installed
    on the running interpreter (so the endpoint never 500s — just
    serves something usable).
    """
    decision_id = decision.get("decision_id", "decision")
    suggested_name = f"{decision_id}.docx"

    # Strip XML-forbidden control characters from every string field so
    # python-docx (lxml under the hood) doesn't raise ValueError on the
    # ~rare scraped decision with stray NULL/form-feed bytes from PDF
    # extraction. See _xml_safe / _xml_safe_decision at top of module.
    decision = _xml_safe_decision(decision)
    if paragraphs:
        paragraphs = [
            {**p, "text": _xml_safe(p.get("text", ""))}
            for p in paragraphs
        ]

    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    except ImportError:
        # Defensive fallback — text version of the same content
        return _render_decision_txt(decision, paragraphs)

    doc = Document()

    # Margins — practitioner-friendly default (2.5 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    # Typography: Times New Roman 12pt, line spacing 1.2 (multiple).
    # Apply to Normal + heading styles so the whole document inherits a
    # single, monochrome serif look — the canonical legal-document style.
    BODY_FONT = "Times New Roman"
    BODY_SIZE = Pt(12)

    def _apply_typography(style):
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.2
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)

    _apply_typography(doc.styles["Normal"])
    for sname in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            _apply_typography(doc.styles[sname])
        except KeyError:
            pass
    # Heading sizes — restrained, monochrome (no color overrides)
    try:
        doc.styles["Heading 1"].font.size = Pt(16)
        doc.styles["Heading 1"].font.bold = True
        doc.styles["Heading 2"].font.size = Pt(13)
        doc.styles["Heading 2"].font.bold = True
    except KeyError:
        pass

    # 1. Citation header (title)
    cite_de = decision.get("citation_string_de") or decision.get("citation_string")
    cite_fr = decision.get("citation_string_fr")
    cite_it = decision.get("citation_string_it")
    title = doc.add_heading(cite_de or _short_title(decision), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 2. Metadata block — court, date, language, joined by commas (no em-dashes,
    # no middots). Plain weight, no italic.
    meta_parts = []
    if decision.get("court_name") or decision.get("court"):
        meta_parts.append(str(decision.get("court_name") or decision.get("court")))
    if decision.get("decision_date"):
        meta_parts.append(str(decision["decision_date"]))
    if decision.get("language"):
        meta_parts.append(str(decision["language"]).upper())
    if meta_parts:
        doc.add_paragraph(", ".join(meta_parts))

    # 3. Source URL
    url_p = doc.add_paragraph()
    url_p.add_run("Quelle: ").bold = True
    url_p.add_run(_decision_url(decision_id))

    # 4. Alternate-language citations
    alt_lines = []
    if cite_fr and cite_fr != cite_de:
        alt_lines.append(f"FR: {cite_fr}")
    if cite_it and cite_it != cite_de:
        alt_lines.append(f"IT: {cite_it}")
    for line in alt_lines:
        doc.add_paragraph(line)

    # 5. Regeste — split multilingual head-note at language boundaries
    regeste = decision.get("regeste") or ""
    regeste_paras = _split_regeste(regeste)
    if regeste_paras:
        doc.add_heading("Regeste", level=2)
        for rp in regeste_paras:
            doc.add_paragraph(rp)

    # 6. Sachverhalt / Erwägungen (structured if available, else fallback)
    if paragraphs:
        doc.add_heading("Erwägungen", level=2)
        try:
            from mcp_server import _e_number_sort_key  # type: ignore
            sorted_paras = sorted(paragraphs, key=lambda p: _e_number_sort_key(
                p.get("e_number") or ""))
        except Exception:
            sorted_paras = paragraphs
        for p in sorted_paras:
            e_num = (p.get("e_number") or "").strip()
            text = (p.get("text") or "").strip()
            if not text:
                continue
            heading = doc.add_paragraph()
            run = heading.add_run(f"E. {e_num}" if e_num else "")
            run.bold = True
            for sub in _split_paragraphs(text):
                doc.add_paragraph(sub)
    else:
        full = (decision.get("full_text") or "").strip()
        if full:
            doc.add_heading("Volltext", level=2)
            for sub in _split_paragraphs(full):
                doc.add_paragraph(sub)

    # 7. Footer disclaimer (plain, no italic, no separator line)
    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Export aus OpenCaseLaw (CC0). Verbindlich ist allein der vom "
        "erlassenden Gericht veröffentlichte Originaltext. Quellen-URL "
        "siehe oben."
    )
    foot.paragraph_format.space_before = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return (
        buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        suggested_name,
    )


# ── PDF (reportlab Platypus) ───────────────────────────────────────

def render_pdf(decision: dict, paragraphs: list[dict] | None = None) -> tuple[bytes, str, str]:
    """Build a PDF with the decision's structured content.

    Same typography as the docx export: Times-Roman 12pt body, leading
    14.4pt (1.2× 12), 2.5 cm margins, A4. Monochrome — no colored
    accents, no horizontal rules. Falls back to plain-text if reportlab
    is not installed (so the endpoint never 500s).
    """
    decision_id = decision.get("decision_id", "decision")
    suggested_name = f"{decision_id}.pdf"

    try:
        from reportlab.lib.pagesizes import A4  # type: ignore
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, KeepTogether,
        )
    except ImportError:
        return _render_decision_txt(decision, paragraphs)

    base = ParagraphStyle(
        "Body",
        fontName="Times-Roman",
        fontSize=12,
        leading=14.4,  # 1.2 × 12pt
        spaceBefore=0,
        spaceAfter=6,
        textColor="#000000",
    )
    h1 = ParagraphStyle(
        "H1", parent=base,
        fontName="Times-Bold", fontSize=16, leading=19.2,
        spaceBefore=0, spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "H2", parent=base,
        fontName="Times-Bold", fontSize=13, leading=15.6,
        spaceBefore=14, spaceAfter=6,
    )
    meta = ParagraphStyle(
        "Meta", parent=base,
        fontSize=11, leading=13.2,
    )
    e_num_style = ParagraphStyle(
        "ErwNum", parent=base,
        fontName="Times-Bold",
        spaceBefore=8, spaceAfter=2,
    )

    def P(text, style=base):
        # reportlab Paragraph parses XML-like markup; escape & < >
        from xml.sax.saxutils import escape as xml_escape
        return Paragraph(xml_escape(text or ""), style)

    story: list = []

    # 1. Citation header
    cite_de = decision.get("citation_string_de") or decision.get("citation_string")
    cite_fr = decision.get("citation_string_fr")
    cite_it = decision.get("citation_string_it")
    story.append(P(cite_de or _short_title(decision), h1))

    # 2. Metadata
    meta_parts = []
    if decision.get("court_name") or decision.get("court"):
        meta_parts.append(str(decision.get("court_name") or decision.get("court")))
    if decision.get("decision_date"):
        meta_parts.append(str(decision["decision_date"]))
    if decision.get("language"):
        meta_parts.append(str(decision["language"]).upper())
    if meta_parts:
        story.append(P(", ".join(meta_parts), meta))

    # 3. Source URL
    story.append(P(f"Quelle: {_decision_url(decision_id)}", meta))

    # 4. Alternate-language citations
    if cite_fr and cite_fr != cite_de:
        story.append(P(f"FR: {cite_fr}", meta))
    if cite_it and cite_it != cite_de:
        story.append(P(f"IT: {cite_it}", meta))

    # 5. Regeste
    regeste = decision.get("regeste") or ""
    regeste_paras = _split_regeste(regeste)
    if regeste_paras:
        story.append(P("Regeste", h2))
        for rp in regeste_paras:
            story.append(P(rp, base))

    # 6. Erwägungen / Volltext
    if paragraphs:
        story.append(P("Erwägungen", h2))
        try:
            from mcp_server import _e_number_sort_key  # type: ignore
            sorted_paras = sorted(paragraphs, key=lambda p: _e_number_sort_key(
                p.get("e_number") or ""))
        except Exception:
            sorted_paras = paragraphs
        for p in sorted_paras:
            e_num = (p.get("e_number") or "").strip()
            text = (p.get("text") or "").strip()
            if not text:
                continue
            if e_num:
                story.append(P(f"E. {e_num}", e_num_style))
            for sub in _split_paragraphs(text):
                story.append(P(sub, base))
    else:
        full = (decision.get("full_text") or "").strip()
        if full:
            story.append(P("Volltext", h2))
            for sub in _split_paragraphs(full):
                story.append(P(sub, base))

    # 7. Footer disclaimer
    story.append(Spacer(1, 12))
    story.append(P(
        "Export aus OpenCaseLaw (CC0). Verbindlich ist allein der vom "
        "erlassenden Gericht veröffentlichte Originaltext. Quellen-URL "
        "siehe oben.", meta,
    ))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        title=cite_de or _short_title(decision),
        author=decision.get("court_name") or decision.get("court") or "OpenCaseLaw",
        subject="Swiss Case Law",
    )
    doc.build(story)
    return buf.getvalue(), "application/pdf", suggested_name


def _render_decision_txt(decision: dict, paragraphs: list[dict] | None) \
        -> tuple[bytes, str, str]:
    """Plain-text fallback when python-docx is unavailable."""
    decision_id = decision.get("decision_id", "decision")
    parts: list[str] = []
    parts.append(decision.get("citation_string_de")
                  or decision.get("citation_string")
                  or _short_title(decision))
    parts.append("=" * len(parts[0]))
    if decision.get("court_name"):
        parts.append(decision["court_name"])
    if decision.get("decision_date"):
        parts.append(decision["decision_date"])
    parts.append(_decision_url(decision_id))
    parts.append("")
    if decision.get("regeste"):
        parts.append("REGESTE")
        parts.append("-------")
        parts.append(decision["regeste"].strip())
        parts.append("")
    if paragraphs:
        parts.append("ERWÄGUNGEN")
        parts.append("----------")
        for p in paragraphs:
            e = (p.get("e_number") or "").strip()
            t = (p.get("text") or "").strip()
            if not t:
                continue
            parts.append(f"E. {e}" if e else "")
            parts.append(t)
            parts.append("")
    elif decision.get("full_text"):
        parts.append("VOLLTEXT")
        parts.append("--------")
        parts.append(decision["full_text"].strip())
    body = "\n".join(parts) + "\n"
    return body.encode("utf-8"), "text/plain; charset=utf-8", \
           f"{decision_id}.txt"


# ── Atom feed (per-court newest decisions) ─────────────────────────

ATOM_NS = 'xmlns="http://www.w3.org/2005/Atom"'


def render_atom_feed(court: str, court_label: str,
                      decisions: Iterable[dict],
                      base_url: str = CANONICAL_BASE) -> tuple[bytes, str, str]:
    """Build an Atom 1.0 feed for the newest decisions of one court.

    `decisions` is an iterable of dicts with at least `decision_id`,
    `decision_date`, `citation_string_de` (or fallback fields), and
    optionally `regeste` for the entry summary. Caller is responsible
    for ordering by date desc and capping the count.
    """
    feed_url = f"{base_url}/atom/{court}.xml"
    site_url = f"{base_url}/courts/{court}"
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    out: list[str] = []
    out.append('<?xml version="1.0" encoding="utf-8"?>')
    out.append(f'<feed {ATOM_NS}>')
    out.append(f"  <title>{escape(court_label)} — neue Entscheide</title>")
    out.append(
        f'  <link rel="self" type="application/atom+xml" href="{escape(feed_url)}"/>'
    )
    out.append(f'  <link rel="alternate" type="text/html" href="{escape(site_url)}"/>')
    out.append(f"  <updated>{now}</updated>")
    out.append(f"  <id>{escape(feed_url)}</id>")
    out.append("  <author><name>OpenCaseLaw</name></author>")
    out.append("  <subtitle>Daily-refreshed feed of newly published decisions. "
                "CC0 — opencaselaw.ch</subtitle>")

    for d in decisions:
        decision_id = d.get("decision_id", "")
        if not decision_id:
            continue
        url = _decision_url(decision_id)
        title = (
            d.get("citation_string_de")
            or d.get("citation_string")
            or _short_title(d)
        )
        date = d.get("decision_date") or now[:10]
        # Atom requires RFC 3339 timestamp
        if re.match(r"^\d{4}-\d{2}-\d{2}$", date):
            date_full = f"{date}T00:00:00Z"
        else:
            date_full = now
        summary = (d.get("regeste") or "")[:500]
        out.append("  <entry>")
        out.append(f"    <id>{escape(url)}</id>")
        out.append(f"    <title>{escape(title)}</title>")
        out.append(
            f'    <link rel="alternate" type="text/html" href="{escape(url)}"/>'
        )
        out.append(f"    <updated>{date_full}</updated>")
        out.append(f"    <published>{date_full}</published>")
        if summary:
            out.append(
                f'    <summary type="text">{escape(summary)}</summary>'
            )
        out.append("  </entry>")
    out.append("</feed>\n")
    body = "\n".join(out)
    return body.encode("utf-8"), "application/atom+xml; charset=utf-8", \
           f"{court}.xml"
